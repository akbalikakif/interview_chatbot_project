"""
reports.py
Mülakat sonuçlarını analiz edip detaylı rapor oluşturan modül
"""

import os
import json
from datetime import datetime
from typing import List, Dict, Any

# PDF ve DOCX için kütüphaneler
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    PDF_AVAILABLE = True
except ImportError:
    print("[UYARI] reportlab kütüphanesi bulunamadı. PDF raporu için: pip install reportlab")
    PDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    print("[UYARI] python-docx kütüphanesi bulunamadı. Word raporu için: pip install python-docx")
    DOCX_AVAILABLE = False

class ReportGenerator:
    """Mülakat raporları oluşturan sınıf"""
    
    def __init__(self):
        self.reports_dir = "reports"
        if not os.path.exists(self.reports_dir):
            os.makedirs(self.reports_dir)
    
    def generate_interview_report(self, interview_history: List[Dict], candidate_name: str = "Aday", chart_paths: Dict[str, str] = None) -> str:
        """Mülakat geçmişinden detaylı rapor oluşturur (grafiklerle)"""
        
        report_data = {
            "candidate_name": candidate_name,
            "interview_date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "total_questions": len(interview_history),
            "scores": self._calculate_scores(interview_history),
            "phases": self._analyze_phases(interview_history),
            "recommendations": self._generate_recommendations(interview_history),
            "detailed_analysis": self._detailed_question_analysis(interview_history)
        }
        
        print("\n" + "="*60)
        print("RAPOR OLUŞTURULUYOR...")
        print("="*60)
        
        # PDF raporu oluştur (grafiklerle)
        pdf_path = None
        if PDF_AVAILABLE:
            try:
                pdf_path = self._create_pdf_report(report_data, chart_paths)
                print(f"[OK] PDF raporu oluşturuldu: {pdf_path}")
            except Exception as e:
                print(f"[HATA] PDF oluşturma hatası: {e}")
        else:
            print("[UYARI] reportlab yüklü değil. PDF raporu oluşturulamadı.")
        
        # DOCX raporu oluştur
        docx_path = None
        if DOCX_AVAILABLE:
            try:
                docx_path = self._create_docx_report(report_data)
                print(f"[OK] Word raporu oluşturuldu: {docx_path}")
            except Exception as e:
                print(f"[HATA] Word oluşturma hatası: {e}")
        else:
            print("[UYARI] python-docx yüklü değil. Word raporu oluşturulamadı.")
        
        print("="*60)
        print("[OK] RAPORLAMA TAMAMLANDI")
        print("="*60)
        
        # Konsol özeti
        summary = self._create_console_summary(report_data)
        return summary
    
    def _calculate_scores(self, history: List[Dict]) -> Dict:
        """Genel puanları hesaplar - LLM ve ses skorlarını ayrı ayrı"""
        if not history:
            return {
                "llm_average": 0, 
                "audio_average": 0, 
                "combined_average": 0,
                "total": 0
            }
        
        # LLM skorları (1-10 arası)
        llm_scores = []
        for h in history:
            if 'analysis' in h and 'score' in h['analysis']:
                llm_scores.append(h['analysis']['score'])
        
        # Ses skorları (0-100 arası, 10'a normalize edilecek)
        audio_scores = []
        for h in history:
            if 'audio_score' in h and h['audio_score']:
                # 0-100 arası skoru 0-10 arası skora çevir
                audio_scores.append(h['audio_score'].get('overall_score', 0) / 10)
        
        # Kombinasyon skoru: LLM %60 + Ses %40
        combined_scores = []
        for h in history:
            llm_score = h.get('analysis', {}).get('score', 0)
            audio_score = 0
            if 'audio_score' in h and h['audio_score']:
                audio_score = h['audio_score'].get('overall_score', 0) / 10
            
            # Eğer ses skoru yoksa sadece LLM skorunu kullan
            if audio_score > 0:
                combined = (llm_score * 0.6) + (audio_score * 0.4)
            else:
                combined = llm_score
            
            combined_scores.append(combined)
        
        result = {
            "llm_average": round(sum(llm_scores) / len(llm_scores), 1) if llm_scores else 0,
            "llm_max": max(llm_scores) if llm_scores else 0,
            "llm_min": min(llm_scores) if llm_scores else 0,
            "audio_average": round(sum(audio_scores) / len(audio_scores), 1) if audio_scores else 0,
            "audio_max": max(audio_scores) if audio_scores else 0,
            "audio_min": min(audio_scores) if audio_scores else 0,
            "combined_average": round(sum(combined_scores) / len(combined_scores), 1) if combined_scores else 0,
            "combined_max": max(combined_scores) if combined_scores else 0,
            "combined_min": min(combined_scores) if combined_scores else 0,
            "total": sum(llm_scores) if llm_scores else 0,
            "count": len(llm_scores),
            "audio_count": len(audio_scores)
        }
        
        # Geriye dönük uyumluluk için
        result["average"] = result["combined_average"]
        result["max"] = result["combined_max"]
        result["min"] = result["combined_min"]
        
        return result
    
    def _analyze_phases(self, history: List[Dict]) -> Dict:
        """Faz bazında analiz yapar"""
        phase_scores = {}
        
        for h in history:
            phase = h.get('kategori', 'bilinmeyen')
            if phase not in phase_scores:
                phase_scores[phase] = []
            
            if 'analysis' in h and 'score' in h['analysis']:
                phase_scores[phase].append(h['analysis']['score'])
        
        # Her faz için ortalama hesapla
        phase_averages = {}
        for phase, scores in phase_scores.items():
            if scores:
                phase_averages[phase] = round(sum(scores) / len(scores), 1)
        
        return phase_averages
    
    def _generate_recommendations(self, history: List[Dict]) -> List[str]:
        """Aday için öneriler oluşturur"""
        recommendations = []
        
        # Genel performans analizi
        scores = [h.get('analysis', {}).get('score', 0) for h in history if 'analysis' in h]
        if not scores:
            return ["Yeterli veri bulunamadı"]
        
        avg_score = sum(scores) / len(scores)
        
        if avg_score >= 8:
            recommendations.append("Mükemmel performans! Teknik bilgi ve iletişim becerileri çok güçlü.")
            recommendations.append("Bu seviyede bir aday için zorlu projeler ve liderlik rolleri önerilir.")
        elif avg_score >= 6:
            recommendations.append("İyi performans. Bazı alanlarda gelişim fırsatları var.")
            recommendations.append("Teknik becerileri güçlendirmek için ek eğitimler önerilir.")
        elif avg_score >= 4:
            recommendations.append("Orta seviye performans. Temel beceriler mevcut ancak gelişim gerekiyor.")
            recommendations.append("Mentorluk ve sürekli eğitim programları önerilir.")
        else:
            recommendations.append("Performans düşük. Temel becerilerin geliştirilmesi gerekiyor.")
            recommendations.append("Kapsamlı eğitim programı ve yakın mentorluk önerilir.")
        
        # Faz bazında öneriler
        phase_scores = self._analyze_phases(history)
        for phase, score in phase_scores.items():
            if score < 5:
                if phase == "kişisel":
                    recommendations.append("Kişisel iletişim becerilerini geliştirmek için sunum ve toplantı pratikleri önerilir.")
                elif phase == "teknik":
                    recommendations.append("Teknik bilgi eksikliği var. İlgili teknolojilerde derinlemesine çalışma önerilir.")
                elif phase == "senaryo":
                    recommendations.append("Problem çözme becerilerini geliştirmek için case study çalışmaları önerilir.")
        
        return recommendations
    
    def _detailed_question_analysis(self, history: List[Dict]) -> List[Dict]:
        """Her soru için detaylı analiz"""
        detailed = []
        
        for i, h in enumerate(history, 1):
            # Zorluk seviyesini hem 'difficulty' hem 'difficulty_level' anahtarlarından kontrol et
            difficulty = h.get('difficulty') or h.get('difficulty_level', 'Bilinmeyen')
            
            # Zorluk seviyesini anlamlı metne çevir
            if isinstance(difficulty, int):
                difficulty_text = {
                    1: "Kolay",
                    2: "Orta",
                    3: "Zor"
                }.get(difficulty, f"Seviye {difficulty}")
            else:
                difficulty_text = str(difficulty) if difficulty != 'Bilinmeyen' else 'Bilinmeyen'
            
            question_analysis = {
                "question_number": i,
                "question": h.get('soru', 'Bilinmeyen soru'),
                "category": h.get('kategori', 'Bilinmeyen'),
                "answer": h.get('answer', 'Cevap bulunamadı'),
                "score": h.get('analysis', {}).get('score', 0),
                "feedback": h.get('analysis', {}).get('feedback', 'Değerlendirme bulunamadı'),
                "difficulty": difficulty_text
            }
            detailed.append(question_analysis)
        
        return detailed
    
    def _create_console_summary(self, report_data: Dict) -> str:
        """Konsol için kısa özet oluşturur"""
        scores = report_data['scores']
        
        summary = f"""
\n{'='*60}
MÜLAKAT ÖZETİ
{'='*60}
Aday: {report_data['candidate_name']}
Tarih: {report_data['interview_date']}
Toplam Soru: {report_data['total_questions']}

GENEL PERFORMANS:
  📊 Genel Skor: {scores['combined_average']}/10
  📝 İçerik Skoru (LLM): {scores['llm_average']}/10
  🎤 Ses Skoru: {scores['audio_average']}/10
  
  En Yüksek: {scores['combined_max']}/10
  En Düşük: {scores['combined_min']}/10
  
  Not: Genel skor = İçerik (%60) + Ses (%40)

FAZ BAZINDA PERFORMANS:
"""
        
        for phase, score in report_data['phases'].items():
            summary += f"  {phase.upper()}: {score}/10\n"
        
        summary += "\nÖNERİLER:\n"
        for i, rec in enumerate(report_data['recommendations'], 1):
            summary += f"  {i}. {rec}\n"
        
        summary += f"\n{'='*60}\n"
        summary += "Detaylı rapor için PDF ve Word dosyalarını kontrol edin.\n"
        summary += f"{'='*60}\n"
        
        return summary
    
    def _create_pdf_report(self, report_data: Dict, chart_paths: Dict[str, str] = None) -> str:
        """PDF formatında rapor oluşturur - Türkçe karakter desteği ve grafiklerle"""
        if not PDF_AVAILABLE:
            return None
        
        pdf_filename = f"interview_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        pdf_path = os.path.join(self.reports_dir, pdf_filename)
        
        # Türkçe karakter desteği için font kaydet
        try:
            # DejaVu Sans fontunu kaydet (Türkçe karakterler için)
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.pdfbase import pdfmetrics
            
            # Windows için font yolları
            font_paths = [
                "C:/Windows/Fonts/DejaVuSans.ttf",
                "C:/Windows/Fonts/Arial.ttf",
                "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",  # Linux
                "/System/Library/Fonts/Supplemental/Arial.ttf"  # macOS
            ]
            
            font_registered = False
            for font_path in font_paths:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont('TurkishFont', font_path))
                        pdfmetrics.registerFont(TTFont('TurkishFont-Bold', font_path))
                        font_registered = True
                        print(f"[OK] Türkçe font kaydedildi: {font_path}")
                        break
                    except Exception as e:
                        continue
            
            if not font_registered:
                print("[UYARI] Türkçe font bulunamadı, varsayılan font kullanılacak")
                turkish_font = 'Helvetica'
                turkish_font_bold = 'Helvetica-Bold'
            else:
                turkish_font = 'TurkishFont'
                turkish_font_bold = 'TurkishFont-Bold'
        except Exception as e:
            print(f"[UYARI] Font kaydı hatası: {e}")
            turkish_font = 'Helvetica'
            turkish_font_bold = 'Helvetica-Bold'
        
        doc = SimpleDocTemplate(pdf_path, pagesize=A4)
        story = []
        styles = getSampleStyleSheet()
        
        # Başlık stili - Türkçe font ile
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName=turkish_font_bold,
            fontSize=24,
            textColor=colors.HexColor('#1a237e'),
            spaceAfter=30,
            alignment=1  # Center
        )
        
        # Alt başlık stili - Türkçe font ile
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontName=turkish_font_bold,
            fontSize=16,
            textColor=colors.HexColor('#283593'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        # Normal metin stili - Türkçe font ile
        normal_style = ParagraphStyle(
            'TurkishNormal',
            parent=styles['Normal'],
            fontName=turkish_font,
            fontSize=11
        )
        
        # Başlık
        story.append(Paragraph("MÜLAKAT DEĞERLENDİRME RAPORU", title_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Genel bilgiler
        info_data = [
            ['Aday:', report_data['candidate_name']],
            ['Tarih:', report_data['interview_date']],
            ['Toplam Soru:', str(report_data['total_questions'])]
        ]
        info_table = Table(info_data, colWidths=[2*inch, 4*inch])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e8eaf6')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), turkish_font_bold),
            ('FONTNAME', (0, 0), (-1, -1), turkish_font),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey)
        ]))
        story.append(info_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Genel Performans
        story.append(Paragraph("GENEL PERFORMANS", heading_style))
        scores = report_data['scores']
        perf_data = [
            ['Metrik', 'Değer'],
            ['Genel Skor (Kombinasyon)', f"{scores['combined_average']}/10"],
            ['İçerik Skoru (LLM)', f"{scores['llm_average']}/10"],
            ['Ses Skoru', f"{scores['audio_average']}/10"],
            ['En Yüksek Puan', f"{scores['combined_max']}/10"],
            ['En Düşük Puan', f"{scores['combined_min']}/10"],
            ['Toplam Puan', str(scores['total'])],
            ['Not', 'Genel skor = İçerik (60%) + Ses (40%)']
        ]
        perf_table = Table(perf_data, colWidths=[3*inch, 3*inch])
        perf_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3f51b5')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), turkish_font_bold),
            ('FONTNAME', (0, 1), (-1, -1), turkish_font),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(perf_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Faz Bazında Performans
        story.append(Paragraph("FAZ BAZINDA PERFORMANS", heading_style))
        phase_data = [['Faz', 'Ortalama Puan']]
        for phase, score in report_data['phases'].items():
            phase_data.append([phase.upper(), f"{score}/10"])
        
        phase_table = Table(phase_data, colWidths=[3*inch, 3*inch])
        phase_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3f51b5')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), turkish_font_bold),
            ('FONTNAME', (0, 1), (-1, -1), turkish_font),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.lightblue),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(phase_table)
        story.append(Spacer(1, 0.3*inch))
        
        # Grafikler ekle (varsa)
        if chart_paths:
            story.append(PageBreak())
            story.append(Paragraph("GELİŞİM GRAFİKLERİ", heading_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Genel gelişim grafiği
            if chart_paths.get('overall_progress') and os.path.exists(chart_paths['overall_progress']):
                story.append(Paragraph("Genel Gelişim", normal_style))
                img = Image(chart_paths['overall_progress'], width=6*inch, height=3*inch)
                story.append(img)
                story.append(Spacer(1, 0.3*inch))
            
            # Faz karşılaştırma
            if chart_paths.get('phase_comparison') and os.path.exists(chart_paths['phase_comparison']):
                story.append(Paragraph("Faz Bazında Performans", normal_style))
                img = Image(chart_paths['phase_comparison'], width=5*inch, height=3*inch)
                story.append(img)
                story.append(Spacer(1, 0.3*inch))
            
            # Radar chart
            if chart_paths.get('radar') and os.path.exists(chart_paths['radar']):
                story.append(Paragraph("Yetenek Haritası", normal_style))
                img = Image(chart_paths['radar'], width=4*inch, height=4*inch)
                story.append(img)
                story.append(Spacer(1, 0.3*inch))
            
            # İyileşme trendi
            if chart_paths.get('improvement_trend') and os.path.exists(chart_paths['improvement_trend']):
                story.append(Paragraph("Kümülatif Gelişim", normal_style))
                img = Image(chart_paths['improvement_trend'], width=6*inch, height=3*inch)
                story.append(img)
                story.append(Spacer(1, 0.3*inch))
        
        # Öneriler
        story.append(PageBreak())
        story.append(Paragraph("ÖNERİLER", heading_style))
        for i, rec in enumerate(report_data['recommendations'], 1):
            story.append(Paragraph(f"{i}. {rec}", normal_style))
            story.append(Spacer(1, 0.1*inch))
        
        story.append(PageBreak())
        
        # Detaylı Soru Analizi
        story.append(Paragraph("DETAYLI SORU ANALİZİ", heading_style))
        
        # Heading3 stili için Türkçe font
        heading3_style = ParagraphStyle(
            'TurkishHeading3',
            parent=styles['Heading3'],
            fontName=turkish_font_bold,
            fontSize=13
        )
        
        for detail in report_data['detailed_analysis']:
            story.append(Paragraph(f"<b>Soru {detail['question_number']} ({detail['category']})</b>", heading3_style))
            story.append(Paragraph(f"<b>Soru:</b> {detail['question']}", normal_style))
            story.append(Paragraph(f"<b>Cevap:</b> {detail['answer'][:200]}{'...' if len(detail['answer']) > 200 else ''}", normal_style))
            story.append(Paragraph(f"<b>Puan:</b> {detail['score']}/10 | <b>Zorluk:</b> {detail['difficulty']}", normal_style))
            story.append(Paragraph(f"<b>Geri Bildirim:</b> {detail['feedback']}", normal_style))
            story.append(Spacer(1, 0.2*inch))
        
        doc.build(story)
        return pdf_path
    
    def _create_docx_report(self, report_data: Dict) -> str:
        """Word (DOCX) formatında rapor oluşturur"""
        if not DOCX_AVAILABLE:
            return None
        
        docx_filename = f"interview_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        docx_path = os.path.join(self.reports_dir, docx_filename)
        
        doc = Document()
        
        # Başlık
        title = doc.add_heading('MÜLAKAT DEĞERLENDİRME RAPORU', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(26, 35, 126)
        
        # Genel Bilgiler
        doc.add_heading('Genel Bilgiler', level=1)
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Light Grid Accent 1'
        table.rows[0].cells[0].text = 'Aday:'
        table.rows[0].cells[1].text = report_data['candidate_name']
        table.rows[1].cells[0].text = 'Tarih:'
        table.rows[1].cells[1].text = report_data['interview_date']
        table.rows[2].cells[0].text = 'Toplam Soru:'
        table.rows[2].cells[1].text = str(report_data['total_questions'])
        
        # Genel Performans
        doc.add_heading('Genel Performans', level=1)
        scores = report_data['scores']
        perf_table = doc.add_table(rows=8, cols=2)
        perf_table.style = 'Light Grid Accent 1'
        perf_table.rows[0].cells[0].text = 'Genel Skor (Kombinasyon)'
        perf_table.rows[0].cells[1].text = f"{scores['combined_average']}/10"
        perf_table.rows[1].cells[0].text = 'İçerik Skoru (LLM)'
        perf_table.rows[1].cells[1].text = f"{scores['llm_average']}/10"
        perf_table.rows[2].cells[0].text = 'Ses Skoru'
        perf_table.rows[2].cells[1].text = f"{scores['audio_average']}/10"
        perf_table.rows[3].cells[0].text = 'En Yüksek Puan'
        perf_table.rows[3].cells[1].text = f"{scores['combined_max']}/10"
        perf_table.rows[4].cells[0].text = 'En Düşük Puan'
        perf_table.rows[4].cells[1].text = f"{scores['combined_min']}/10"
        perf_table.rows[5].cells[0].text = 'Toplam Puan'
        perf_table.rows[5].cells[1].text = str(scores['total'])
        perf_table.rows[6].cells[0].text = 'Değerlendirilen Soru'
        perf_table.rows[6].cells[1].text = str(scores['count'])
        perf_table.rows[7].cells[0].text = 'Not'
        perf_table.rows[7].cells[1].text = 'Genel skor = İçerik (60%) + Ses (40%)'
        
        # Faz Bazında Performans
        doc.add_heading('Faz Bazında Performans', level=1)
        phase_table = doc.add_table(rows=len(report_data['phases'])+1, cols=2)
        phase_table.style = 'Light Grid Accent 1'
        phase_table.rows[0].cells[0].text = 'Faz'
        phase_table.rows[0].cells[1].text = 'Ortalama Puan'
        for i, (phase, score) in enumerate(report_data['phases'].items(), 1):
            phase_table.rows[i].cells[0].text = phase.upper()
            phase_table.rows[i].cells[1].text = f"{score}/10"
        
        # Öneriler
        doc.add_heading('Öneriler', level=1)
        for i, rec in enumerate(report_data['recommendations'], 1):
            doc.add_paragraph(f"{i}. {rec}", style='List Number')
        
        # Detaylı Soru Analizi
        doc.add_page_break()
        doc.add_heading('Detaylı Soru Analizi', level=1)
        for detail in report_data['detailed_analysis']:
            doc.add_heading(f"Soru {detail['question_number']} ({detail['category']})", level=2)
            doc.add_paragraph(f"Soru: {detail['question']}")
            doc.add_paragraph(f"Cevap: {detail['answer'][:300]}{'...' if len(detail['answer']) > 300 else ''}")
            doc.add_paragraph(f"Puan: {detail['score']}/10 | Zorluk: {detail['difficulty']}")
            doc.add_paragraph(f"Geri Bildirim: {detail['feedback']}")
            doc.add_paragraph('')  # Boşluk
        
        doc.save(docx_path)
        return docx_path

def generate_final_report(interview_handler) -> str:
    """InterviewHandler'dan rapor oluşturur"""
    generator = ReportGenerator()
    return generator.generate_interview_report(interview_handler.history)